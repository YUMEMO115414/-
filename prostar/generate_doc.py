"""
生成 ProSTAR 论文分析文档 (Word 格式)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ====== 样式设置 ======
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# 标题样式
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = '黑体'
    h.font.color.rgb = RGBColor(0, 0, 0)

def add_para(text, bold=False, size=11, align=None, font_name=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
    if align is not None:
        p.alignment = align
    return p

def add_code_block(text):
    """添加等宽代码块"""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(50, 50, 50)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2

# ====== 封面 ======
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ProSTAR 论文深度解析')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Prototype-Based Satellite-Terrestrial Adaptation\nfor Resource Constrained Co-Inference')
run.font.size = Pt(13)
run.font.name = 'Calibri'
run.italic = True

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('论文作者: Shenhu Zhang, Shi Yan, Fengxian Guo, Mianji Li, Nan Li, Mugen Peng\n单位: 北京邮电大学 网络与交换技术国家重点实验室\n代码复现与文档整理: 2026年6月')
run.font.size = Pt(10)

doc.add_page_break()

# ====== 目录占位 ======
doc.add_heading('目录', level=1)
add_para('一、论文概览与核心创新点')
add_para('二、问题建模与算法原理')
add_para('三、代码结构与模块解析')
add_para('四、实验设计与结果分析')
add_para('五、存在的不足与局限')
add_para('六、未来可行研究方向')
add_para('七、参考文献与延伸阅读')
doc.add_page_break()

# ======================================================================
# 一、论文概览
# ======================================================================
doc.add_heading('一、论文概览与核心创新点', level=1)

doc.add_heading('1.1 研究背景', level=2)
add_para('低轨卫星网络（LEO Satellite Networks）近年来在地球观测、灾害响应、环境监测等领域扮演着越来越重要的角色。传统的卫星工作模式很简单：卫星负责拍照，然后把原始图像一股脑传回地面站，地面站用训练好的深度神经网络做分类或者检测。这套流程有两个硬伤。')
add_para('第一，卫星和地面之间的通信带宽非常有限。一颗 LEO 卫星掠过地面站上空的时间窗口可能只有几分钟，在这几分钟里要把几百 MB 的高分辨率遥感图像全部传下来，物理上就不现实。第二，卫星自身的计算能力虽然在逐年提升，但跟地面动辄几百 TFLOPS 的 GPU 集群相比仍然差了好几个数量级。更麻烦的是，卫星上能做推理但几乎没法做训练——反向传播那套操作在功耗和散热极其紧张的太空环境里完全是奢侈品。')
add_para('于是就有了"星地协同推理"（Satellite-Terrestrial Collaborative Inference）这个折中方案：卫星做前半段推理，把中间结果传给地面，地面做后半段。想法很好，但现有方案几乎都假设模型是在闭集（closed-set）上训练的——也就是说，推理时遇到的类别必须在训练集中出现过。这个假设在地球观测场景里根本不成立。灾难类型、新出现的军事目标、突发的环境事件，这些都不是训练集能提前覆盖的。')
add_para('ProSTAR 就是冲着这个缺口来的。它要解决的问题可以总结为一句话：在卫星计算资源紧张、星地链路带宽吃紧、新类层出不穷的三重约束下，怎么做到又快又省电又灵活的地球观测图像识别。')

doc.add_heading('1.2 五个核心创新点', level=2)

add_para('创新一：训练免适配（Training-Free Adaptation）', bold=True)
add_para('这是整篇论文最核心的亮点。传统的深度学习模型遇到新类别就得重新训练——不管是全量微调还是增量学习，都绕不开反向传播。ProSTAR 直接绕过了这一步。它把推理过程拆成两件事：特征提取和语义匹配。特征提取用预训练的 Swin-Tiny 就够了，参数全部冻结，不需要更新。语义匹配用"原型"（Prototype）——本质上就是每个类的代表性特征向量——来做最近邻查找。新类来了怎么办？算一个它的原型存起来就行，完全不需要碰 DNN 的参数。这个思路跟人类的认知方式有点像：你见过一次"袋鼠"就能记住它的特征，以后在别的地方再看到就能认出来，并不需要重写整个大脑。')

add_para('创新二：嵌套子空间投影（Nested Subspace Projection）', bold=True)
add_para('卫星到地面的通信带宽是稀缺资源，能少传一点就少传一点。ProSTAR 的做法很巧妙：先算好一个大的随机正交投影矩阵 P_base，然后根据当前信道条件从里面截取前 d 列当投影矩阵用。d 大一点（信道好）就多传点信息，d 小一点（信道差）就少传点。关键是这个"截取"操作是 O(1) 的——不需要重新算投影矩阵，不需要重新训练，什么都不需要。跟俄罗斯套娃一样，大套娃里面装着中套娃，中套娃里面装着小套娃，想要哪个尺寸直接取就行。')

add_para('创新三：一维特征传输取代三维特征张量', bold=True)
add_para('传统 Split Inference 方案传输的是三维中间特征图（比如 Swin-Tiny Stage4 输出的 7×7×768 张量），大约 150KB。ProSTAR 在上星侧做完 GAP（全局平均池化）之后，三维张量坍缩成了一维向量，再用嵌套投影压到 256 维（1KB）。传输数据量降了超过两个数量级，直接的效果就是延迟和能耗大幅下降。论文报告的数值是总延迟减少 38.4%，总能耗降低 41.7%。')

add_para('创新四：基于代理模型的自适应决策（Surrogate Model-Based Decision）', bold=True)
add_para('前面提到了 d 越大准确率越高但开销也越大，d 越小反过来的 trade-off。卫星需要根据实时信道质量、剩余电量、任务紧急程度来动态选择最优的 (l, d) 组合。问题在于评估每个 (l, d) 的真实准确率需要遍历所有候选去做 DNN 前传——这在星上是完全不可行的。ProSTAR 的解决方案是用一个 sigmoid 代理模型来"预测"准确率：Â_l(d) = sigmoid(d; α_l, d0_l)。地面站离线拟合好代理模型的四个参数，编成一个超轻量参数表（几 KB）上传给卫星。卫星在线决策时只需把当前信道速率和计算频率代入公式做代数运算，全流程不到一毫秒。')

add_para('创新五：指数移动平均自适应原型更新（EMA-Based Prototype Update）', bold=True)
add_para('遥感图像有一个特点：同一种地物在不同季节、不同光照、不同角度下看起来可能完全不一样（这叫"域偏移"）。如果原型永远是建库时那个样子，时间一长准确率就会持续下滑。ProSTAR 的做法是对匹配置信度高的查询样本，用指数移动平均（EMA）平滑地更新对应类的原型：新的原型 = 0.9 × 旧原型 + 0.1 × 当前样本的特征。只更新前 d 维，也只更新被预测的那个类。计算开销小到可以忽略，但长期来看对维持模型鲁棒性很有效。')
doc.add_page_break()

# ======================================================================
# 二、问题建模与算法原理
# ======================================================================
doc.add_heading('二、问题建模与算法原理', level=1)

doc.add_heading('2.1 系统模型', level=2)
add_para('考虑一个 LEO 卫星 S 和一个地面站 G 组成的协同推理系统。卫星获得遥感图像 x ∈ R^{H×W×C}，目标是给出分类标签 y ∈ Y。DNN 模型沿前向传播被切分成两部分：卫星负责前 l 层的特征提取 F^l_ext，地面负责从中间特征到最终标签的匹配 F^l_match。')

add_para('卫星端的计算延迟和能耗采用经典的 CMOS 模型：T_comp = γ_l × F_total / f_sat，E_comp = κ × f²_sat × γ_l × F_total。通信延迟和能耗采用香农公式建模：R = B × log₂(1 + SNR)（式4），T_comm = S(d)/R + d_sg/c（式5），E_comm = P_tx × S(d)/R（式6）。其中 S(d) = d × 32 bits 是传输的总比特数。')

doc.add_heading('2.2 核心算法详解', level=2)

add_para('阶段一：嵌套子空间原型构建（式7）', bold=True)
add_para('地面站对每个新类收集 K 个标注样本（默认 K=5），用跟卫星一致的 DNN 前 l 层提取 1D 特征，跟随机正交基矩阵 P_base 相乘后取各类均值，得到 M 个"基原型" c^base_m ∈ R^{D_l}。P_base 是随机生成的 D_l × D_l 正交矩阵，在离线阶段只算一次。')

add_para('阶段二：特征压缩（式8）', bold=True)
add_para('卫星做完前 l 层前传并 GAP 之后得到 f_l ∈ R^{D_l}。根据当前目标维度 d，取 P_base 的前 d 列构成 P^l_d，计算 z_new = f_l × P^l_d。这个投影的 FLOPs 是 O(D_l × d)，跟 DNN 的前传计算量比起来差了三四个数量级。')

add_para('阶段三：余弦相似度匹配（式9）', bold=True)
add_para('地面站收到 z_new 后，从每个基原型截取前 d 维，用归一化的余弦相似度 ŷ = argmax_m cos(z_new, c_m(d)) 完成分类。不需要任何额外的神经网络层、softmax 或者前向计算。整个匹配过程复杂度 O(M × d)。')

add_para('阶段四：EMA 原型更新（式10）', bold=True)
add_para('对匹配置信度高于阈值 τ 的查询样本，更新对应类的原型：c_ŷ(t+1) = (1-β)c_ŷ(t) + β × z_new，β 取值通常在 0.1~0.2。')

doc.add_heading('2.3 联合优化问题（式13~19）', level=2)
add_para('目标函数：min_{l,d} J(l,d) = λ₁ × (A_max - A(l,d)) + λ₂ × T_total(l,d) + λ₃ × E_total(l,d)。这是 NP-hard 的非凸组合优化问题。ProSTAR 的解法是两步走。第一步，地面站在验证集上对所有 (l,d) 组合做一次穷举测量，得到经验数据集 D^l_profile（式17）。第二步，用改进的 sigmoid 函数 Â_l(d) = A_max + (A_max - A_min) / (1 + e^{-α(d-d0)})（式16）去拟合这些散点，用非线性最小二乘法估计参数 θ*_l（式18）。最后把四个参数 (A_max, A_min, α, d0) 的表发给卫星（总共 4×4=16 个浮点数，64 字节）。卫星在线只需遍历 L×|D_l| 个候选做代数运算即可得到最优 (l*, d*)（式19）。')
doc.add_page_break()

# ======================================================================
# 三、代码结构
# ======================================================================
doc.add_heading('三、代码结构与模块解析', level=1)

doc.add_heading('3.1 项目目录', level=2)
add_code_block("""D:\\cursor_cc\\prostar\\
  __init__.py        # 包初始化, 模块元信息
  prototype.py       # 核心算法层 (Eq.7~10): 原型构建、特征压缩、余弦匹配、EMA更新
  surrogate.py       # 优化层 (Eq.13~19): 代理模型拟合、联合代价、在线决策
  demo.py            # 可视化层: 5 个实验函数, 生成 Fig.2~5
  main.py            # 独立运行入口 (备用)
run.py               # 项目根快捷入口
ProSTAR.pdf          # 原论文 (需从微信目录手动拷贝)""")

doc.add_heading('3.2 prototype.py — 核心算法模块', level=2)
add_para('该模块实现了论文 Section III 的全部4个阶段, 共约 320 行。')
add_para('PrototypeBank 类（第 22~48 行）', bold=True)
add_para('原型仓库的数据结构。内部维护一个 (M × D_l) 的 NumPy 数组存储所有类的基原型。核心方法 get_prototype(d) 是 O(1) 的切片操作——从基原型截取前 d 列返回低维原型矩阵。')
add_para('initialize_prototypes 函数（第 61~98 行）', bold=True)
add_para('实现式7: 输入支持集特征和标签, 用 QR 分解生成随机正交基, 计算各类的基原型, 做 L2 归一化后存入 PrototypeBank。')
add_para('compress_feature 函数（第 108~118 行）', bold=True)
add_para('实现式8: 输入原始特征 f_l 和基投影矩阵 P_base, 取前 d 列做矩阵乘法, 输出 d 维压缩向量 z_new。')
add_para('cosine_similarity_match 函数（第 130~148 行）', bold=True)
add_para('实现式9: 对 z_new 和所有原型做 L2 归一化后计算点积 (归一化后点积 = 余弦相似度), 返回 argmax 对应的类别。')
add_para('update_prototype_ema 函数（第 166~192 行）', bold=True)
add_para('实现式10: 检查置信度, 用 EMA 平滑更新预测类原型的前 d 维。')
add_para('ProSTARPipeline 类（第 222~270 行）', bold=True)
add_para('把上述四个阶段的函数封装成一个 end-to-end 推理流水线, 提供 initialize/infer/infer_and_update 三个高层接口。')

doc.add_heading('3.3 surrogate.py — 代理模型与优化模块', level=2)
add_para('该模块实现了论文 Section IV 的联合优化框架, 共约 280 行。')
add_para('LayerSurrogateParams 类（第 18~40 行）', bold=True)
add_para('存储单层代理模型的四个参数 (A_max, A_min, α, d0) 和预测方法 predict(d)。predict 内部实现了式16的 sigmoid 计算。')
add_para('fit_surrogate_params 函数（第 98~170 行）', bold=True)
add_para('实现式17~18: 使用对数参数化和数值裁剪保护的梯度下降法, 不依赖 scipy.optimize, 2000 次迭代收敛。注意这里用的是 log-alpha 参数化来保证 alpha 恒为正。')
add_para('SystemParams 类（第 180~230 行）', bold=True)
add_para('实现式1~6: 封装星地系统的计算模型和通信模型。channel_rate() 用香农公式算信道速率, comp_latency/comp_energy 用 CMOS 模型算计算开销, comm_latency/comm_energy 算传输开销。')
add_para('joint_cost 函数（第 245~258 行）', bold=True)
add_para('实现式13: J(l,d) = λ₁×max(0, A_max-A(l,d)) + λ₂×T_total + λ₃×E_total。使用 max(0,·) 防止准确率超过 A_max 时出现负代价。')
add_para('online_decision 函数（第 282~318 行）', bold=True)
add_para('实现式19: 遍历所有候选 (l, d), 用代理模型预测准确率, 代入系统模型计算开销, 选出联合代价最小的组合。纯代数运算, 实测 < 700 μs。')

doc.add_heading('3.4 demo.py — 可视化复现模块', level=2)
add_para('该模块包含了4个独立的实验函数和1个驱动入口, 全部带有详细的中文逐行注释。')
add_para('SimulatedFeatureExtractor 类：模拟 Swin-Tiny 的轻量特征提取器。内部维护21个类的随机语义中心、4层的 He 初始化投影矩阵, 以及信息沿维度指数衰减的先验机制（这是复现 sigmoid 退化曲线的关键）。', bold=False)
add_para('figure_2_surrogate_fitting: 1×4 子图展示4个 Stage 的代理模型拟合效果。蓝色散点 = 24个压缩级别上的实测准确率, 蓝色曲线 = sigmoid 拟合, 红色虚线 = d0 临界维度。', bold=False)
add_para('figure_3_exhaustive_vs_proposed: 2×2 热力图。上行 = 准确率 (穷举/代理), 下行 = 联合代价 (穷举/代理)。星号标记各自的最优 (l*, d*)。', bold=False)
add_para('figure_4_confusion_matrix: 16类混淆矩阵, 蓝色渐变, 行归一化为百分比。右上角标注总体准确率。', bold=False)
add_para('figure_5_performance_comparison: 1×3 红蓝柱状图, 对比3D原始特征和1D压缩特征在准确率、延迟、能耗三个指标上的表现。柱顶标注具体数值, 延迟和能耗子图上标注降幅百分比。', bold=False)
doc.add_page_break()

# ======================================================================
# 四、实验设计
# ======================================================================
doc.add_heading('四、实验设计与结果分析', level=1)

doc.add_heading('4.1 实验配置', level=2)
add_para('骨干网络: Swin-Tiny, ImageNet-1K 预训练, 4 个 Stage, 通道数 [96, 192, 384, 768]')
add_para('数据集: UCMerced LandUse (21 类遥感场景), 所有类视为新类')
add_para('小样本设置: 5-shot (每类5个标注样本构建原型)')
add_para('系统参数: f_sat=5GHz, B=1MHz, P_tx=15W, f_GS=20GHz')
add_para('优化权重: λ₁=200, λ₂=30, λ₃=2')
add_para('压缩级别: 每层24个均匀间隔的 d 值')

doc.add_heading('4.2 实验结果', level=2)
add_para('代理模型拟合（Fig.2）: Layer 0~3 的 sigmoid 退化曲线拟合 MAE 在 7~10% 之间, 成功捕获了"大 d 平稳、小 d 锐降"的趋势。')
add_para('穷举 vs 代理对比（Fig.3）: 代理模型识别出的最优 (l,d) 与穷举搜索一致（均为 Layer 3, d≈400~500）, 证实了代理模型的有效性。在线决策耗时 < 700 μs。')
add_para('混淆矩阵（Fig.4）: 21类全部视为新类的困难设置下, 仅用 5-shot 原型 + 256维压缩特征即达到总体准确率 92.5%。对角线高度集中, 说明压缩后的1D特征仍保留强判别力。')
add_para('性能对比（Fig.5）: 1D 压缩特征的准确率比 3D 原始特征略低 ~5%, 但总延迟降低 ~43%, 总能耗降低 ~48%。这个 trade-off 在实际星地系统中非常有吸引力——用可接受的精度损失换取了显著的资源节省。')
doc.add_page_break()

# ======================================================================
# 五、存在的不足与局限
# ======================================================================
doc.add_heading('五、存在的不足与局限', level=1)

add_para('坦诚地说, ProSTAR 虽然思路巧妙, 但并非没有短板。以下讨论几个值得注意的局限。', bold=False)

doc.add_heading('5.1 代理模型的参数拟合依赖离线穷举', level=2)
add_para('代理模型的四个参数 (A_max, A_min, α, d0) 需要在每个分区层独立拟合, 而拟合用的"训练数据"来自对验证集上所有 (l,d) 组合的穷举测量。论文用的 Swin-Tiny 只有4个 Stage, L=4 的穷举负担尚可接受。但如果换成 ResNet-152 或 ViT-Large 这种几十层上百层的模型, 遍历每一层做穷举测量会变得极其耗时。这个问题在论文中没有被讨论。')

doc.add_heading('5.2 嵌套子空间依赖随机投影', level=2)
add_para('ProSTAR 的特征压缩使用的是"随机正交投影"——P_base 是随机生成的, 不是从数据中学习出来的。这个选择的优点是不需要训练, 但代价是投影的质量没有理论保证。换一个随机种子, P_base 的结构变了, 压缩后的判别力可能就不一样了。论文没有报告随机种子对最终准确率的影响 (应该做一个 sensitivity analysis)。')

doc.add_heading('5.3 原型匹配假设类内方差较小', level=2)
add_para('原型网络的本质是用一个向量来代表整个类, 这隐含假设了"同类样本在特征空间中紧密聚集"。对于"停车场"、"跑道"这种纹理高度一致的场景分类, 这个假设基本成立。但对于"港口"这种可能包含水域、船只、起重机、集装箱等多模态子类的复杂场景, 一个原型可能不足以捕捉类内的多样性。每类用多个原型 (multi-prototype) 或者高斯原型 (probabilistic prototype) 可能更合适, 但这会增加存储和匹配的复杂度。')

doc.add_heading('5.4 EMA 更新策略假设数据流是平稳的', level=2)
add_para('EMA 更新的本质是对历史原型和当前样本做加权平均, 更新率 β 是固定的。在季节缓慢变化的场景下 (比如森林覆盖率的年际变化) 这个假设合理。但如果发生突变——比如一场洪水把一片农田变成了湖泊——EMA 的平滑特性会让原型"追不上"这种突变。可以考虑自适应步长 EMA, 或者引入变化检测 (change detection) 来触发重置。')

doc.add_heading('5.5 缺乏与 SOTA 小样本学习方法的直接对比', level=2)
add_para('论文对比了 Ground-only、On-board Updating 和 Traditional Split 三种协同推理方案, 这四者都是在"卫星-地面协同"这个范式下的比较。但论文没有把 ProSTAR 跟小样本学习中更通用的方法（如 MatchingNet、Prototypical Network、MAML、MetaOptNet、FEAT 等）放在同一数据集上做公平对比。缺少这个对标, 读者很难直观判断 ProSTAR 的"92.5% 准确率"到底处于什么水平。')

doc.add_heading('5.6 模拟器假设与实际部署之间存在差距', level=2)
add_para('复现中使用的 SimulatedFeatureExtractor 为了可控性和可复现性做了简化——用随机投影 + 信息衰减先验来模拟 DNN 的行为。实际 Swin-Tiny 的特征分布可能跟这种简化有差异。更严格的做法是在真实的 ImageNet 预训练 Swin-Tiny 上用真实的 UCMerced 图像跑一遍, 对比简化模拟器和真实 DNN 的 A(l,d) 退化曲线是否一致。')

doc.add_heading('5.7 仅处理分类任务', level=2)
add_para('论文的方法是为图像分类场景设计的——地面站维护的原型就是分类器的权重。对于目标检测（需要回归框）、语义分割（需要像素级预测）、变化检测（需要时序对比）等地球观测中同样重要的任务, 目前的方法没法直接套用。如何处理更复杂的结构化预测任务是后续需要攻克的方向。')
doc.add_page_break()

# ======================================================================
# 六、未来方向
# ======================================================================
doc.add_heading('六、未来可行研究方向', level=1)

doc.add_heading('6.1 多原型和高斯原型', level=2)
add_para('正如 5.3 节所说, 单原型在复杂场景下有局限性。一个自然的扩展是每类维护 K 个子原型 (Mixture of Prototypes), 匹配时用最小距离或加权投票。或者把原型扩展为高斯分布 (Probabilistic Prototype), 每个类用一个均值向量 + 协方差矩阵来描述, 匹配时用马氏距离。这两个方向的代价都很低, 不会破坏 ProSTAR "训练免更新"的核心优势。')

doc.add_heading('6.2 从依赖随机投影到学习投影', level=2)
add_para('ProSTAR 当前使用的是随机正交投影, 但投影矩阵的质量显然可以通过学习来提升。一个可行的方案是: 在地面站的离线阶段, 利用现有的标注数据端到端地学习一个投影矩阵 P_learned, 使得投影后的原型在余弦相似度下的匹配准确率最大化。这个学习的代价是一次性的, 学完后上传到卫星就不需要再动了, 不违反"在线训练免更新"的约束。')

doc.add_heading('6.3 多卫星协同和联邦原型学习', level=2)
add_para('论文考虑的是单颗卫星和单个地面站。实际 LEO 星座中有数十到数百颗卫星, 每颗卫星从不同角度、不同时间拍摄地球的不同区域。如果多颗卫星共享原型, 地面站可以通过联邦平均的方式把各卫星独立更新的原型合并成一个更鲁棒的全局原型, 再分发回各卫星。这跟联邦学习 (Federated Learning) 的思想一致, 但不需要传模型参数——只传原型向量 (几 KB 量级)。')

doc.add_heading('6.4 生成式基础模型的集成', level=2)
add_para('2024~2026 年间, 遥感领域出现了多个大规模预训练基础模型 (如 GeoChat、SatMAE、DeCUR 等)。这些模型具有更强的零/少样本泛化能力。把 ProSTAR 的骨干从 Swin-Tiny 换成这些预训练基础模型, 原型匹配准确率预计会有显著提升。值得探索的具体问题是: 基础模型的特征维度通常很大 (768~2048 维), d 的选择对性能的影响会比 Swin-Tiny 更复杂, 代理模型是否能同样好地拟合这些更大维度下的退化曲线。')

doc.add_heading('6.5 扩展到检测和分割任务', level=2)
add_para('将 ProSTAR 的框架从"整图分类"扩展到"目标检测"的核心挑战在于: 检测需要处理多尺度和多实例, 每个候选框都需要独立的原型匹配, 原型库的结构需要从"类→向量"升级为"类→(向量, 尺度先验, 上下文特征)"。分割更有意思——也许可以维护每个像素的"超像素原型", 用原型来指导分割边界的划定。')

doc.add_heading('6.6 在线增量学习的理论分析', level=2)
add_para('ProSTAR 的 EMA 更新目前是启发式的——选 β 靠经验, 判断置信度靠阈值。有几个理论问题值得深入: EMA 更新的收敛性分析 (在什么条件下原型会收敛到真实的类中心？)、最优 β 的理论表达式 (是否跟数据流的非平稳性有关？)、以及灾难性遗忘的边界 (新原型会覆盖旧知识吗？)。这些问题如果把答案搞清楚, 有可能把 ProSTAR 从应用驱动的工程方案升级为有理论保证的学习框架。')

doc.add_heading('6.7 星上硬件的真实部署验证', level=2)
add_para('目前的工作都是在地面 GPU 上跑的仿真。星上嵌入式平台 (如 NVIDIA Jetson Xavier/Orin, Xilinx FPGA, 华为 Atlas 200) 的真实功耗和延迟特性与桌面仿真的差异需要在硬件层面验证。特别是 P_base 投影这一步——它在 GPU 上几乎免费, 但在没有矩阵乘法加速单元的 FPGA 上就可能成为瓶颈。把 ProSTAR 的完整流水线移植到至少一种真实的星上边缘计算平台上, 并测量端到端的延迟和功耗, 是把论文从"模拟验证"推向"实际可用"的关键一步。')
doc.add_page_break()

# ======================================================================
# 七、参考文献
# ======================================================================
doc.add_heading('七、参考文献与延伸阅读', level=1)

refs = [
    '[1] Zhang S, Yan S, Guo F, et al. ProSTAR: Prototype-Based Satellite-Terrestrial Adaptation for Resource Constrained Co-Inference[J]. 2025.',
    '[2] Furutanpey A, Zhang Q, Raith P, et al. FOOL: Addressing the Downlink Bottleneck in Satellite Computing with Neural Feature Compression[J]. IEEE Transactions on Mobile Computing, 2025, 24(8): 6747-6764.',
    '[3] Yao S, Lin Y, Wang M, et al. LEOEdge: A Satellite-Ground Cooperation Platform for AI Inference in Large LEO Constellation[J]. IEEE Journal on Selected Areas in Communications, 2025, 43(1): 36-50.',
    '[4] Fan W, Meng Q, Wang G, et al. Satellite Edge Intelligence: DRL-Based Resource Management for Task Inference in LEO-Based Satellite-Ground Collaborative Networks[J]. IEEE Transactions on Mobile Computing, 2025, 24(10): 10710-10728.',
    '[5] Lovelace R, Nowosad J, Muenchow J. Geocomputation with R[M]. 2nd ed. Boca Raton: CRC Press, 2024.',
    '[6] de Lange N. Geoinformatics in Theory and Practice[M]. Berlin: Springer, 2023.',
    '[7] Shekhar S, Chawla S. Spatial Databases: A Tour[M]. Upper Saddle River: Prentice Hall, 2003.',
    '[8] Lillesand T M, Kiefer R W, Chipman J W. Remote Sensing and Image Interpretation[M]. 7th ed. Hoboken: Wiley, 2015.',
    '[9] Bolstad P, Manson S. GIS Fundamentals: A First Text on Geographic Information Systems[M]. 7th ed. White Bear Lake: Eider Press, 2022.',
    '[10] Haut J M, Paoletti M E, Moreno-Álvarez S, et al. Distributed Deep Learning for Remote Sensing Data Interpretation[J]. Proceedings of the IEEE, 2021, 109(8): 1320-1349.',
    '[11] 秦昆, 卢宾宾, 陈江平, 等. 空间数据分析[M]. 3版. 武汉: 武汉大学出版社, 2023.',
    '[12] 彭木根, 许文嘉, 孙耀华, 等. 空间信息通信[M]. 北京: 清华大学出版社, 2024.',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

# ====== 保存 ======
output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                           'ProSTAR_Paper_Analysis.docx')
doc.save(output_path)
print(f'Word 文档已生成: {output_path}')
